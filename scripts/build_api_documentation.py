from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("Campus_Directory_API_Documentation.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
      shd = OxmlElement("w:shd")
      tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def style_run(run, bold=False, size=None, color=None, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text.strip())
    style_run(run, size=8.5, font="Consolas")
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F6F9")
    p_pr.append(shd)
    return p


def add_note(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "E8EEF5")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    style_run(r, bold=True, color="1F4D78")
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)


def add_endpoint_table(doc, rows):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_table_width(table, [0.8, 2.35, 1.0, 2.35])
    headers = ["Method", "Endpoint", "Auth", "Fungsi"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, "E8EEF5")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        style_run(run, bold=True, color="1F4D78")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if i in (0, 2):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(value)
            style_run(run, size=9.5)
    doc.add_paragraph()


def add_payload(doc, title, method, endpoint, auth, body=None, curl=None, notes=None):
    h = doc.add_heading(title, level=3)
    h.paragraph_format.keep_with_next = True
    meta = doc.add_table(rows=4 if notes else 3, cols=2)
    meta.style = "Table Grid"
    set_table_width(meta, [1.3, 5.2])
    items = [("Method", method), ("Endpoint", endpoint), ("Auth", auth)]
    if notes:
        items.append(("Catatan", notes))
    for idx, (k, v) in enumerate(items):
        left = meta.cell(idx, 0)
        right = meta.cell(idx, 1)
        set_cell_shading(left, "F2F4F7")
        r = left.paragraphs[0].add_run(k)
        style_run(r, bold=True, color="1F4D78")
        rr = right.paragraphs[0].add_run(v)
        style_run(rr, size=9.5)
    if body:
        doc.add_paragraph("Payload / Body:")
        add_code(doc, body)
    if curl:
        doc.add_paragraph("Contoh cURL:")
        add_code(doc, curl)


def configure_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        s = styles[name]
        s.font.name = "Calibri"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        s.font.size = Pt(size)
        s.font.color.rgb = RGBColor.from_string(color)
        s.font.bold = True
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)


def build():
    doc = Document()
    configure_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    r = title.add_run("Campus Directory API")
    style_run(r, bold=True, size=24, color="0B2545")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = subtitle.add_run("Dokumentasi endpoint dan payload testing")
    style_run(rs, size=12, color="555555")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = meta.add_run(f"Backend: Node.js + Express + PostgreSQL | Dibuat: {date.today().isoformat()}")
    style_run(rm, size=9.5, color="555555")

    add_note(
        doc,
        "Ringkasan cepat",
        "Gunakan BASE_URL sesuai environment: http://localhost:3000 untuk lokal, atau URL deploy Railway/Render/Neon. Endpoint admin membutuhkan token user dengan role admin.",
    )

    doc.add_heading("1. Setup Testing", level=1)
    doc.add_paragraph("Variable yang disarankan untuk Postman/Thunder Client:")
    add_code(doc, """
BASE_URL=http://localhost:3000
TOKEN=<isi_dari_response_login>
ADMIN_TOKEN=<token_admin>
PLACE_ID=1
CATEGORY_ID=1
REVIEW_ID=1
""")
    doc.add_paragraph("Header umum:")
    add_code(doc, """
Content-Type: application/json
Authorization: Bearer {{TOKEN}}
""")
    add_note(
        doc,
        "Akun admin default",
        "Email: admin@campus.ac.id | Password: Admin@123. Ganti password setelah login pertama jika dipakai di server publik.",
    )

    doc.add_heading("2. Format Response", level=1)
    doc.add_paragraph("Success:")
    add_code(doc, '{"status":"success","message":"OK","data":{}}')
    doc.add_paragraph("Paginated:")
    add_code(doc, '{"status":"success","data":[],"meta":{"total":30,"page":1,"limit":20,"pages":2}}')
    doc.add_paragraph("Error:")
    add_code(doc, '{"status":"error","message":"Deskripsi error","errors":[]}')

    doc.add_heading("3. Ringkasan Endpoint", level=1)
    add_endpoint_table(doc, [
        ("GET", "/health", "-", "Cek status server"),
        ("GET", "/api", "-", "Daftar endpoint API"),
        ("POST", "/api/auth/register", "-", "Registrasi user"),
        ("POST", "/api/auth/login", "-", "Login dan ambil JWT"),
        ("GET", "/api/auth/me", "Bearer", "Profil user login"),
        ("PUT", "/api/auth/me", "Bearer", "Update profil/avatar"),
        ("PUT", "/api/auth/change-password", "Bearer", "Ganti password"),
        ("GET", "/api/categories", "-", "Daftar kategori"),
        ("GET", "/api/categories/:id", "-", "Detail kategori"),
        ("POST", "/api/categories", "Admin", "Tambah kategori"),
        ("PUT", "/api/categories/:id", "Admin", "Update kategori"),
        ("DELETE", "/api/categories/:id", "Admin", "Hapus kategori"),
        ("GET", "/api/places", "Optional", "Daftar/filter tempat"),
        ("GET", "/api/places/nearby", "-", "Tempat terdekat dari GPS"),
        ("GET", "/api/places/:id", "Optional", "Detail tempat"),
        ("POST", "/api/places", "Admin", "Tambah tempat"),
        ("PUT", "/api/places/:id", "Admin", "Update tempat"),
        ("DELETE", "/api/places/:id", "Admin", "Soft delete tempat"),
        ("GET", "/api/places/:id/photos", "-", "Daftar foto tempat"),
        ("POST", "/api/places/:id/photos", "Admin", "Tambah foto"),
        ("GET", "/api/places/:id/reviews", "-", "Daftar review tempat"),
        ("GET", "/api/places/:id/reviews/summary", "-", "Ringkasan rating"),
        ("POST", "/api/places/:id/reviews", "Bearer", "Tambah review"),
        ("PUT", "/api/reviews/:id", "Owner/Admin", "Update review"),
        ("DELETE", "/api/reviews/:id", "Owner/Admin", "Hapus review"),
        ("GET", "/api/favorites", "Bearer", "Daftar favorit user"),
        ("POST", "/api/favorites/:placeId", "Bearer", "Tambah favorit"),
        ("DELETE", "/api/favorites/:placeId", "Bearer", "Hapus favorit"),
        ("GET", "/api/favorites/check/:placeId", "Bearer", "Cek favorit"),
        ("GET", "/api/admin/stats", "Admin", "Dashboard statistik"),
        ("GET", "/api/admin/users", "Admin", "Daftar user"),
        ("PUT", "/api/admin/users/:id/role", "Admin", "Ubah role"),
        ("PUT", "/api/admin/users/:id/toggle", "Admin", "Aktif/nonaktif user"),
        ("GET", "/api/admin/places", "Admin", "Semua tempat"),
        ("PUT", "/api/admin/places/:id/verify", "Admin", "Toggle verifikasi tempat"),
    ])

    doc.add_heading("4. Auth", level=1)
    add_payload(doc, "Register", "POST", "{{BASE_URL}}/api/auth/register", "-", """
{
  "name": "Budi Santoso",
  "email": "budi@example.com",
  "password": "secret123"
}
""", """
curl -X POST "{{BASE_URL}}/api/auth/register" ^
  -H "Content-Type: application/json" ^
  -d "{\\"name\\":\\"Budi Santoso\\",\\"email\\":\\"budi@example.com\\",\\"password\\":\\"secret123\\"}"
""")
    add_payload(doc, "Login", "POST", "{{BASE_URL}}/api/auth/login", "-", """
{
  "email": "admin@campus.ac.id",
  "password": "Admin@123"
}
""", """
curl -X POST "{{BASE_URL}}/api/auth/login" ^
  -H "Content-Type: application/json" ^
  -d "{\\"email\\":\\"admin@campus.ac.id\\",\\"password\\":\\"Admin@123\\"}"
""", "Simpan field data.token sebagai ADMIN_TOKEN atau TOKEN.")
    add_payload(doc, "Get My Profile", "GET", "{{BASE_URL}}/api/auth/me", "Bearer {{TOKEN}}")
    add_payload(doc, "Update Profile", "PUT", "{{BASE_URL}}/api/auth/me", "Bearer {{TOKEN}}", """
{
  "name": "Budi Updated"
}
""", notes="Endpoint juga menerima form-data field avatar untuk upload gambar.")
    add_payload(doc, "Change Password", "PUT", "{{BASE_URL}}/api/auth/change-password", "Bearer {{TOKEN}}", """
{
  "old_password": "secret123",
  "new_password": "secret456"
}
""")

    doc.add_heading("5. Categories", level=1)
    add_payload(doc, "Get Categories", "GET", "{{BASE_URL}}/api/categories", "-")
    add_payload(doc, "Get Category Detail", "GET", "{{BASE_URL}}/api/categories/{{CATEGORY_ID}}", "-")
    add_payload(doc, "Create Category", "POST", "{{BASE_URL}}/api/categories", "Bearer {{ADMIN_TOKEN}}", """
{
  "name": "Lab Komputer",
  "icon": "monitor",
  "color": "#2563EB"
}
""")
    add_payload(doc, "Update Category", "PUT", "{{BASE_URL}}/api/categories/{{CATEGORY_ID}}", "Bearer {{ADMIN_TOKEN}}", """
{
  "name": "Laboratorium",
  "icon": "monitor",
  "color": "#1D4ED8"
}
""")
    add_payload(doc, "Delete Category", "DELETE", "{{BASE_URL}}/api/categories/{{CATEGORY_ID}}", "Bearer {{ADMIN_TOKEN}}", notes="Akan gagal jika kategori masih dipakai oleh tempat aktif.")

    doc.add_heading("6. Places", level=1)
    doc.add_paragraph("Query params GET /api/places yang didukung:")
    add_endpoint_table(doc, [
        ("Param", "category", "-", "ID kategori atau nama kategori, contoh cafe"),
        ("Param", "search", "-", "Cari nama/deskripsi/alamat"),
        ("Param", "lat,lng", "-", "Koordinat user untuk menghitung distance_km"),
        ("Param", "radius", "-", "Filter radius kilometer, butuh lat/lng"),
        ("Param", "tags", "-", "CSV tags, contoh wifi,ac"),
        ("Param", "sort", "-", "name, rating, atau distance"),
        ("Param", "page,limit", "-", "Paginasi"),
    ])
    add_payload(doc, "List Places", "GET", "{{BASE_URL}}/api/places?category=cafe&search=kopi&page=1&limit=10", "-")
    add_payload(doc, "List Places With Distance", "GET", "{{BASE_URL}}/api/places?lat=-7.556&lng=112.228&radius=2&sort=distance", "-")
    add_payload(doc, "Nearby Places", "GET", "{{BASE_URL}}/api/places/nearby?lat=-7.556&lng=112.228&radius=1&category=cafe&limit=20", "-", notes="Dipakai untuk fitur GPS dan marker terdekat.")
    add_payload(doc, "Place Detail", "GET", "{{BASE_URL}}/api/places/{{PLACE_ID}}?lat=-7.556&lng=112.228", "-")
    add_payload(doc, "Create Place", "POST", "{{BASE_URL}}/api/places", "Bearer {{ADMIN_TOKEN}}", """
{
  "category_id": 1,
  "name": "Kafe Demo",
  "address": "Jl. Kampus Demo No. 1",
  "lat": -7.5561,
  "lng": 112.2281,
  "description": "Tempat testing API.",
  "phone": "081234567890",
  "open_hours": "Senin-Jumat 08.00-22.00",
  "website": "https://example.com",
  "price_range": "Rp10K-25K",
  "tags": ["wifi", "ac", "colokan"]
}
""", notes="Bisa juga form-data dengan field photo untuk upload gambar.")
    add_payload(doc, "Update Place", "PUT", "{{BASE_URL}}/api/places/{{PLACE_ID}}", "Bearer {{ADMIN_TOKEN}}", """
{
  "category_id": 1,
  "name": "Kafe Demo Updated",
  "address": "Jl. Kampus Demo No. 2",
  "lat": -7.5562,
  "lng": 112.2282,
  "description": "Deskripsi sudah diperbarui.",
  "phone": "081234567890",
  "open_hours": "Setiap hari 08.00-23.00",
  "website": "https://example.com",
  "price_range": "Rp15K-30K",
  "tags": "wifi,ac,parkir",
  "is_active": true,
  "is_verified": true
}
""")
    add_payload(doc, "Delete Place", "DELETE", "{{BASE_URL}}/api/places/{{PLACE_ID}}", "Bearer {{ADMIN_TOKEN}}", notes="Soft delete: is_active menjadi false.")
    add_payload(doc, "Get Place Photos", "GET", "{{BASE_URL}}/api/places/{{PLACE_ID}}/photos", "-")
    add_payload(doc, "Add Place Photo", "POST", "{{BASE_URL}}/api/places/{{PLACE_ID}}/photos", "Bearer {{ADMIN_TOKEN}}", notes="Gunakan form-data: key photo, type File, value gambar jpg/jpeg/png/webp.")

    doc.add_heading("7. Reviews", level=1)
    add_payload(doc, "Get Place Reviews", "GET", "{{BASE_URL}}/api/places/{{PLACE_ID}}/reviews?page=1&limit=20", "-")
    add_payload(doc, "Get Review Summary", "GET", "{{BASE_URL}}/api/places/{{PLACE_ID}}/reviews/summary", "-")
    add_payload(doc, "Create Review", "POST", "{{BASE_URL}}/api/places/{{PLACE_ID}}/reviews", "Bearer {{TOKEN}}", """
{
  "rating": 5,
  "comment": "Tempat nyaman, wifi stabil, cocok untuk nugas."
}
""", notes="Satu user hanya bisa memberi satu review per tempat.")
    add_payload(doc, "Update Review", "PUT", "{{BASE_URL}}/api/reviews/{{REVIEW_ID}}", "Bearer {{TOKEN}}", """
{
  "rating": 4,
  "comment": "Update komentar review."
}
""")
    add_payload(doc, "Delete Review", "DELETE", "{{BASE_URL}}/api/reviews/{{REVIEW_ID}}", "Bearer {{TOKEN}}", notes="Owner review atau admin boleh menghapus.")

    doc.add_heading("8. Favorites", level=1)
    add_payload(doc, "Get My Favorites", "GET", "{{BASE_URL}}/api/favorites", "Bearer {{TOKEN}}")
    add_payload(doc, "Add Favorite", "POST", "{{BASE_URL}}/api/favorites/{{PLACE_ID}}", "Bearer {{TOKEN}}")
    add_payload(doc, "Remove Favorite", "DELETE", "{{BASE_URL}}/api/favorites/{{PLACE_ID}}", "Bearer {{TOKEN}}")
    add_payload(doc, "Check Favorite", "GET", "{{BASE_URL}}/api/favorites/check/{{PLACE_ID}}", "Bearer {{TOKEN}}")

    doc.add_heading("9. Admin", level=1)
    add_payload(doc, "Dashboard Stats", "GET", "{{BASE_URL}}/api/admin/stats", "Bearer {{ADMIN_TOKEN}}")
    add_payload(doc, "List Users", "GET", "{{BASE_URL}}/api/admin/users?page=1&limit=20&search=budi", "Bearer {{ADMIN_TOKEN}}")
    add_payload(doc, "Set User Role", "PUT", "{{BASE_URL}}/api/admin/users/{{USER_ID}}/role", "Bearer {{ADMIN_TOKEN}}", """
{
  "role": "admin"
}
""")
    add_payload(doc, "Toggle User Active", "PUT", "{{BASE_URL}}/api/admin/users/{{USER_ID}}/toggle", "Bearer {{ADMIN_TOKEN}}")
    add_payload(doc, "List All Places", "GET", "{{BASE_URL}}/api/admin/places?page=1&limit=20&active=true", "Bearer {{ADMIN_TOKEN}}")
    add_payload(doc, "Toggle Place Verification", "PUT", "{{BASE_URL}}/api/admin/places/{{PLACE_ID}}/verify", "Bearer {{ADMIN_TOKEN}}")

    doc.add_heading("10. Checklist Testing Demo", level=1)
    checks = [
        "Jalankan npm install.",
        "Buat .env dari .env.example dan isi DB_PASSWORD atau DATABASE_URL.",
        "Jalankan npm run db:migrate untuk membuat tabel dan seed data.",
        "Jalankan npm start lalu buka {{BASE_URL}}/health.",
        "Login admin dan simpan token.",
        "Test GET /api/categories, GET /api/places, dan GET /api/places/1.",
        "Test nearby dengan lat/lng untuk membuktikan fitur GPS.",
        "Test create/update place memakai token admin.",
        "Test review/favorite memakai token user biasa.",
        "Ganti BASE_URL ke URL deployment saat demo dari HP.",
    ]
    for item in checks:
        p = doc.add_paragraph(style=None)
        p.style = doc.styles["List Bullet"]
        p.add_run(item)

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_heading("Lampiran: Form-data Upload", level=1)
    add_endpoint_table(doc, [
        ("Endpoint", "/api/auth/me", "Bearer", "avatar: File"),
        ("Endpoint", "/api/places", "Admin", "photo: File + field tempat"),
        ("Endpoint", "/api/places/:id", "Admin", "photo: File + field tempat"),
        ("Endpoint", "/api/places/:id/photos", "Admin", "photo: File"),
    ])
    doc.add_paragraph("Format gambar yang diizinkan: jpg, jpeg, png, webp. Ukuran maksimal mengikuti MAX_FILE_SIZE_MB di .env, default 5 MB.")

    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT.resolve())
